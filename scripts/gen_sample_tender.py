"""Generate a synthetic bilingual (EN/AR) tender document set for Phase 0 E2E verification.

Produces:
  - sample_tender/ITT_Tender_Document.docx  (Instructions to Tenderers + project metadata + requirements)
  - sample_tender/BOQ_Bill_of_Quantities.xlsx (BOQ line items across trades)

These exercise the v1 pipeline end-to-end: parsing (Docling DOCX + openpyxl XLSX),
bilingual/Arabic handling, summary extraction (13 fields), and checklist extraction
(technical/commercial/legal/HSE/submission/eligibility).

Run:  .venv/Scripts/python.exe scripts/gen_sample_tender.py
Requires: python-docx, openpyxl
"""

from __future__ import annotations

from pathlib import Path

OUT = Path(__file__).resolve().parent.parent / "sample_tender"
OUT.mkdir(parents=True, exist_ok=True)


def build_docx() -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    def h(text: str, level: int = 1):
        doc.add_heading(text, level=level)

    def p(text: str):
        para = doc.add_paragraph(text)
        for run in para.runs:
            run.font.size = Pt(11)
        return para

    # ── Title page ──────────────────────────────────────────────
    h("INSTRUCTIONS TO TENDERERS (ITT)", 0)
    p("Project: New Cairo Medical Center - Main Hospital Building")
    p("اسم المشروع: مركز القاهرة الجديدة الطبي - مبنى المستشفى الرئيسي")
    p("Tender Reference No.: NCMC-2026-TND-014")
    p("")

    # ── Section 1: Project Information ───────────────────────────
    h("1. Project Information / معلومات المشروع", 1)
    p("Project Owner / Employer: New Urban Communities Authority (NUCA).")
    p("صاحب العمل: هيئة المجتمعات العمرانية الجديدة.")
    p("Main Contractor (Client of Bidder): Hassan Allam Construction.")
    p("Location: Plot 7, Medical District, New Cairo, Cairo Governorate, Egypt.")
    p("الموقع: قطعة 7، الحي الطبي، القاهرة الجديدة، محافظة القاهرة، مصر.")
    p(
        "Project Description: Design and construction of a 12-storey, 350-bed general "
        "hospital with two basement levels for parking and MEP plant, a helipad, and "
        "ancillary external works. Gross built-up area approximately 48,500 square metres."
    )
    p(
        "Scope of Work (S.O.W): The Works comprise all civil, structural, architectural, "
        "mechanical, electrical, plumbing (MEP), medical gas, fire-fighting, HVAC, "
        "low-current systems, landscaping, and infrastructure connections required for a "
        "fully operational hospital, delivered on a turnkey basis."
    )
    p("نطاق العمل: تشمل الأعمال جميع الأعمال المدنية والإنشائية والمعمارية والكهروميكانيكية والتكييف ومكافحة الحريق والأعمال الخارجية اللازمة لتشغيل المستشفى بالكامل.")
    p("Project Duration: 30 months from the date of the Letter of Commencement.")
    p("مدة المشروع: 30 شهرًا من تاريخ خطاب البدء.")

    # ── Section 2: Key Dates ────────────────────────────────────
    h("2. Key Dates / التواريخ الهامة", 1)
    p("ITT Release Date: 01 March 2026.")
    p("Pre-Bid (Pre-Tender) Meeting and Site Visit: 15 March 2026 at 10:00 AM, on site.")
    p("اجتماع ما قبل المناقصة وزيارة الموقع: 15 مارس 2026 الساعة العاشرة صباحًا.")
    p("Deadline for Clarification Requests: 25 March 2026.")
    p("Tender Submission Deadline: 20 April 2026 at 12:00 noon (Cairo local time).")
    p("الموعد النهائي لتقديم العطاءات: 20 أبريل 2026 الساعة الثانية عشرة ظهرًا.")
    p("Bid Validity Period: 120 days from the submission deadline.")
    p("مدة سريان العطاء: 120 يومًا من الموعد النهائي للتقديم.")
    p("Estimated Award Date: 01 June 2026.")

    # ── Section 3: Commercial & Contract Terms ──────────────────
    h("3. Commercial and Contract Terms / الشروط التجارية والتعاقدية", 1)
    p("Contract Type: Lump Sum (fixed price) contract based on FIDIC Red Book 1999, as amended.")
    p("نوع العقد: عقد بقيمة إجمالية مقطوعة وفقًا لعقد الفيديك الأحمر 1999.")
    p("Governing Law: The laws of the Arab Republic of Egypt.")
    p("Tender Bond (Bid Bond): EGP 5,000,000 (five million Egyptian Pounds), valid for 150 days.")
    p("التأمين الابتدائي: 5,000,000 جنيه مصري ساري لمدة 150 يومًا.")
    p("Tender Documents Fee: EGP 25,000, non-refundable.")
    p("Performance Bond: 10% of the Contract Price, in the form of an unconditional bank guarantee.")
    p("ضمان حسن التنفيذ: 10% من قيمة العقد.")
    p("Advance Payment: 15% of the Contract Price against an equivalent advance payment guarantee.")
    p("الدفعة المقدمة: 15% من قيمة العقد مقابل خطاب ضمان.")
    p("Retention: 5% of each interim payment, capped at 5% of the Contract Price.")
    p("نسبة المحتجز: 5% من كل دفعة مرحلية بحد أقصى 5% من قيمة العقد.")
    p("Payment Cycle: Monthly interim payment certificates; payment within 60 days of certification.")
    p("Liquidated Damages: 0.1% of the Contract Price per day of delay, capped at 10%.")
    p("Defects Liability / Warranty Period: 365 days from Taking-Over.")
    p("Sustainability Target: LEED Gold certification (USGBC) is mandatory for the Works.")
    p("الاستدامة: شهادة ليد ذهبية إلزامية.")

    # ── Section 4: Stakeholders ─────────────────────────────────
    h("4. Consultants and Stakeholders / الاستشاريون والأطراف", 1)
    p("Design Consultant / Engineer: Dar Al-Handasah (Shair and Partners).")
    p("Project Management Consultant (PMC): Hill International.")
    p("Cost Consultant / Quantity Surveyor: AECOM.")
    p("Supervision Consultant: Dar Al-Handasah.")

    # ── Section 5: Technical Requirements ───────────────────────
    h("5. Technical Requirements / المتطلبات الفنية", 1)
    p("5.1 All structural concrete shall comply with ECP 203-2018 and ACI 318. Minimum characteristic strength C35/45 for vertical elements.")
    p("5.2 Reinforcement steel shall be high-tensile deformed bars grade B500B to BS 4449.")
    p("5.3 The Contractor shall submit a detailed BIM model (LOD 350) using Autodesk Revit for all disciplines.")
    p("5.4 HVAC systems shall be designed to ASHRAE 170 for healthcare ventilation, with HEPA filtration in operating theatres.")
    p("5.5 Medical gas systems shall comply with HTM 02-01 and NFPA 99.")
    p("5.6 All materials shall be submitted for the Engineer's approval via material submittals prior to procurement.")
    p("5.7 Mock-ups shall be provided for typical patient room, facade unit, and operating theatre.")

    # ── Section 6: Commercial Requirements ──────────────────────
    h("6. Commercial Requirements / المتطلبات التجارية", 1)
    p("6.1 Tenderers shall price every BOQ item; unpriced items shall be deemed included in other rates.")
    p("6.2 Rates shall be in Egyptian Pounds (EGP) and shall be fixed and not subject to escalation.")
    p("6.3 All applicable taxes including 14% VAT shall be shown separately.")
    p("6.4 The Tenderer shall submit a priced preliminaries (general items) schedule.")

    # ── Section 7: Legal Requirements ───────────────────────────
    h("7. Legal Requirements / المتطلبات القانونية", 1)
    p("7.1 The Tenderer must be a legal entity registered in Egypt with a valid commercial register and tax card.")
    p("7.2 Disputes shall be resolved by arbitration under the Cairo Regional Centre for International Commercial Arbitration (CRCICA).")
    p("7.3 The Tenderer shall comply with all applicable Egyptian labour and social insurance laws.")

    # ── Section 8: HSE Requirements ─────────────────────────────
    h("8. Health, Safety and Environment (HSE) / الصحة والسلامة والبيئة", 1)
    p("8.1 The Contractor shall submit a project-specific HSE Plan within 14 days of award.")
    p("8.2 A full-time certified Safety Officer (NEBOSH IGC) shall be present on site at all times.")
    p("8.3 The Contractor shall maintain an Environmental Management Plan compliant with Egyptian Law 4/1994.")

    # ── Section 9: Submission Documents ─────────────────────────
    h("9. Documents to be Submitted / المستندات المطلوب تقديمها", 1)
    p("9.1 Completed and signed Form of Tender.")
    p("9.2 Tender Bond (bid bond) in the required amount and form.")
    p("9.3 Priced Bill of Quantities (hard copy and Excel soft copy).")
    p("9.4 Construction Programme (Primavera P6) and method statements.")
    p("9.5 Company profile, audited financial statements for the last three years, and similar project references.")
    p("9.6 Valid commercial registration, tax card, and contractor classification certificate.")

    # ── Section 10: Eligibility / Pre-Qualification ─────────────
    h("10. Eligibility and Pre-Qualification Criteria / معايير التأهيل", 1)
    p("10.1 The Tenderer shall hold a Federation of Egyptian Construction Contractors classification of First Grade in Buildings.")
    p("10.2 Minimum average annual turnover of EGP 800,000,000 over the last three financial years.")
    p("10.3 The Tenderer shall have completed at least two hospital or healthcare projects of comparable size in the last ten years.")
    p("10.4 The Tenderer shall demonstrate access to a credit line of at least EGP 200,000,000.")

    out = OUT / "ITT_Tender_Document.docx"
    doc.save(out)
    return out


def build_boq() -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    wb = Workbook()
    ws = wb.active
    ws.title = "BOQ"

    headers = ["Item No.", "Description", "Unit", "Quantity", "Trade"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)

    rows = [
        ["1.1", "Excavation in all types of soil to required levels and disposal", "m3", 18500, "Civil"],
        ["1.2", "Plain concrete blinding C15/20 under foundations, 100mm thick", "m2", 4200, "Civil"],
        ["2.1", "Reinforced concrete C35/45 in raft foundation including formwork", "m3", 6800, "Concrete"],
        ["2.2", "Reinforced concrete C35/45 in columns and shear walls", "m3", 5400, "Concrete"],
        ["2.3", "High-tensile reinforcement steel B500B supply, cut, bend and fix", "ton", 4900, "Concrete"],
        ["3.1", "Supply and install LV power distribution boards (MDB/SMDB)", "no", 86, "Electrical"],
        ["3.2", "Supply and lay XLPE LV cables including trays and containment", "m", 42000, "Electrical"],
        ["3.3", "Supply and install medical IT isolated power panels for OTs", "no", 14, "Electrical"],
        ["4.1", "Supply and install air-cooled chillers, 500 TR each", "no", 6, "HVAC"],
        ["4.2", "Supply and install AHUs with HEPA filtration for operating theatres", "no", 22, "HVAC"],
        ["4.3", "GI ductwork supply and installation including insulation", "m2", 31000, "HVAC"],
        ["5.1", "Medical gas copper pipework (oxygen, vacuum, medical air) to HTM 02-01", "m", 9800, "Plumbing"],
        ["5.2", "Supply and install sanitary fixtures (WCs, basins, scrub sinks)", "no", 1250, "Plumbing"],
        ["6.1", "Fire-fighting wet riser, sprinkler network and pumps to NFPA 13", "lot", 1, "Fire Fighting"],
    ]
    for r in rows:
        ws.append(r)

    # widen columns
    widths = {"A": 10, "B": 64, "C": 8, "D": 12, "E": 16}
    for col, w in widths.items():
        ws.column_dimensions[col].width = w

    out = OUT / "BOQ_Bill_of_Quantities.xlsx"
    wb.save(out)
    return out


if __name__ == "__main__":
    docx_path = build_docx()
    print(f"[ok] wrote {docx_path}")
    boq_path = build_boq()
    print(f"[ok] wrote {boq_path}")
    print(f"[ok] sample tender set in: {OUT}")
