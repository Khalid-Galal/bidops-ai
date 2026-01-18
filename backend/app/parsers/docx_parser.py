"""DOCX document parser using python-docx."""

import time
from pathlib import Path
from typing import Optional

from app.parsers.base import BaseParser, ParsedContent, ParserRegistry


@ParserRegistry.register
class DocxParser(BaseParser):
    """Parser for Microsoft Word documents (.docx)."""

    supported_extensions = [".docx"]
    supported_mimetypes = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    async def parse(self, file_path: str) -> ParsedContent:
        """Parse a DOCX document.

        Args:
            file_path: Path to DOCX file

        Returns:
            ParsedContent with extracted text and metadata
        """
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError

        self.validate_file(file_path)
        start_time = time.time()

        try:
            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Check for heading styles
                    if para.style and para.style.name.startswith("Heading"):
                        paragraphs.append(f"\n## {text}\n")
                    else:
                        paragraphs.append(text)

            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                tables.append({
                    "index": table_idx,
                    "data": table_data,
                })

                # Add table text to paragraphs
                for row in table_data:
                    paragraphs.append(" | ".join(row))

            full_text = "\n".join(paragraphs)

            # Extract metadata
            metadata = await self.extract_metadata(file_path)

            processing_time = int((time.time() - start_time) * 1000)

            return ParsedContent(
                text=full_text,
                metadata=metadata,
                tables=tables if tables else None,
                processing_time_ms=processing_time,
            )

        except PackageNotFoundError:
            raise ValueError(f"Invalid or corrupted DOCX file: {file_path}")
        except Exception as e:
            raise Exception(f"Failed to parse DOCX: {str(e)}") from e

    async def extract_metadata(self, file_path: str) -> dict:
        """Extract metadata from DOCX.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary of metadata
        """
        from docx import Document

        try:
            doc = Document(file_path)
            props = doc.core_properties

            return {
                "title": props.title or None,
                "author": props.author or None,
                "subject": props.subject or None,
                "keywords": props.keywords or None,
                "created": props.created.isoformat() if props.created else None,
                "modified": props.modified.isoformat() if props.modified else None,
                "last_modified_by": props.last_modified_by or None,
                "revision": props.revision or None,
                "category": props.category or None,
                "file_size": Path(file_path).stat().st_size,
            }
        except Exception:
            return {
                "file_size": Path(file_path).stat().st_size,
            }
